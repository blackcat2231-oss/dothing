import streamlit as st
import google.generativeai as genai
from PIL import Image
import pandas as pd
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import io
import time

# --- 1. 網頁基本設定 ---
st.set_page_config(page_title="篤行幼兒園評量系統", layout="wide", page_icon="🌱")

st.markdown("""
    <style>
    .main {background-color: #f9f9f9;}
    .stHeader {color: #2c3e50;}
    th {
        white-space: normal !important;
        min-width: 120px;
        vertical-align: top !important;
        background-color: #f0f2f6 !important;
    }
    td {text-align: center !important; vertical-align: middle !important;}
    td:last-child {text-align: left !important;}
    </style>
    """, unsafe_allow_html=True)

# --- 2. 側邊欄與 API 設定 ---
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/2231/2231649.png", width=100)
    st.title("🌱 篤行幼兒園")
    st.subheader("評量系統 v1.8 (全能整合版)")
    
    if "GEMINI_API_KEY" in st.secrets:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        st.success("API 連線狀態：🟢 線上")
    else:
        st.error("API Key 未設定")
        st.stop()
        
    st.markdown("---")
    menu = st.radio("功能選單", ["📝 批次上傳與辨識", "📄 產生整合評量報告"])

# --- 3. 核心功能函式 ---

def get_gemini_model():
    """尋找最強模型"""
    model_list = []
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                model_list.append(m.name)
        priority_keywords = ['gemini-3', 'gemini-2.5', 'pro', 'flash']
        for keyword in priority_keywords:
            for name in model_list:
                if keyword in name and ('1.5' in name if keyword == 'flash' else True):
                    return genai.GenerativeModel(name), name
    except:
        pass
    return genai.GenerativeModel('gemini-1.5-flash'), 'fallback-flash'

def analyze_image(image):
    """AI 辨識核心：加入學習區判斷"""
    model, model_name = get_gemini_model()
    
    prompt = """
    你是一位專業的資料輸入員。請分析這張幼兒園評量表。
    
    【任務一：判斷學習區】
    請閱讀表頭，判斷這張表屬於哪個學習區？(例如：語文區、數學區、美勞區、積木區、體能區...)。
    請將結果放入 "area" 欄位。

    【任務二：讀取表頭指標】
    請讀取表格最上方、位於「幼兒姓名」與「備註」中間的那 4 個欄位標題文字。
    
    【任務三：讀取幼兒資料】
    請依序讀取每一列幼兒的資料。
    
    **關於「備註」的指示：**
    1. 將格子內所有文字合併。
    2. 保留換行或編號 (①, 1.)。

    **關於「分數」的判斷：**
    - 圈選 1 -> "A"
    - 圈選 2 -> "R"
    - 圈選 3 -> "D"
    - 圈選 4 -> "N"

    【輸出格式】
    回傳 JSON：
    {
      "area": "語文區",
      "headers": ["指標1文字", "指標2文字", "指標3文字", "指標4文字"],
      "students": [
        {"name": "幼兒一", "scores": ["A", "R", "A", "R"], "note": "備註內容..."},
        ...
      ]
    }
    """
    
    config = genai.types.GenerationConfig(temperature=0.0, response_mime_type="application/json")
    try:
        response = model.generate_content([prompt, image], generation_config=config)
        return json.loads(response.text)
    except:
        return None

def generate_teacher_comments(student_name, records):
    """
    AI 寫手核心：
    將該位幼兒的所有區域資料打包，請 AI 寫出像範例一樣的「綜合觀察」與「居家建議」。
    """
    model, _ = get_gemini_model()
    
    # 將資料轉為文字描述給 AI 看
    data_summary = f"幼兒姓名：{student_name}\n"
    for r in records:
        data_summary += f"--- {r['area']} ---\n"
        data_summary += f"指標與成績：{r['details']}\n" # details 包含指標名稱與分數
        data_summary += f"老師原始備註：{r['note']}\n"
    
    prompt = f"""
    你是一位資深的幼兒園園長與教育專家。請根據以下這位幼兒在不同學習區的評量數據與老師備註，撰寫一份給家長的綜合評語。
    
    【幼兒資料】
    {data_summary}
    
    【撰寫目標】
    請模仿以下風格，撰寫兩個段落：
    
    1. **【老師的觀察】**：
       - 綜合所有區域的表現，找出孩子的亮點（哪裡表現好/A級）。
       - 溫柔地指出需要協助的地方（哪裡是D或N），並將其描述為「發展中的珍貴階段」。
       - 語氣要溫暖、正向、專業。
       - 如果原始備註有具體事件（如「恐龍王國」），請務必寫進去，讓故事更生動。
       
    2. **【居家互動小撇步】**：
       - 針對孩子較弱的項目（R/D/N），給家長具體、簡單、可在家進行的遊戲或互動建議。
       - 如果孩子都很好，則建議如何延伸挑戰。
    
    【輸出格式】
    請直接回傳 JSON，不要 markdown：
    {{
        "observation": "這裡寫老師的觀察段落...",
        "suggestion": "這裡寫居家互動小撇步..."
    }}
    """
    
    config = genai.types.GenerationConfig(temperature=0.7, response_mime_type="application/json") # 稍微提高溫度讓文筆更好
    try:
        response = model.generate_content(prompt, generation_config=config)
        return json.loads(response.text)
    except:
        return {"observation": "AI 撰寫中...", "suggestion": "建議親師保持密切聯繫。"}

def create_integrated_word(grouped_data):
    """產生整合版 Word 報告"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft JhengHei'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    
    # 進度顯示
    progress_text = "正在撰寫報告..."
    my_bar = st.progress(0, text=progress_text)
    total_students = len(grouped_data)
    
    for idx, (name, records) in enumerate(grouped_data.items()):
        # 更新進度
        my_bar.progress((idx + 1) / total_students, text=f"正在為 {name} 撰寫評語 ({idx+1}/{total_students})...")
        
        if idx > 0: doc.add_page_break()
        
        # 1. 標題
        head = doc.add_heading('篤行非營利幼兒園', 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph('幼兒學習區個別評量報告')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].bold = True
        sub.runs[0].font.size = Pt(14)
        
        doc.add_paragraph(f"幼兒姓名：{name} \t\t\t 日期：2026年___月___日")
        
        # 2. 呼叫 AI 進行綜合寫作 (這是最花時間的一步)
        ai_comments = generate_teacher_comments(name, records)
        
        # 3. 建立大表格 (包含各區)
        doc.add_paragraph("■ 各區學習指標明細").runs[0].bold = True
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "學習指標"
        hdr[1].text = "評量結果"
        
        # 填入各區資料
        for record in records:
            # 加入區域標題列 (例如：語文區)
            row_area = table.add_row().cells
            row_area[0].merge(row_area[1])
            run = row_area[0].paragraphs[0].add_run(f"【{record['area']}】")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 102, 204) # 藍色標題
            row_area[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 加入該區的指標
            # record['details'] 是一個 list: [{"idx": "能閱讀...", "score": "A"}, ...]
            for item in record['details']:
                row = table.add_row().cells
                row[0].text = item['idx']
                row[1].text = item['score']
                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")
        
        # 4. 寫入 AI 生成的評語
        doc.add_heading('親師交流與建議', level=2)
        
        obs_title = doc.add_paragraph("【老師的觀察】")
        obs_title.runs[0].bold = True
        doc.add_paragraph(ai_comments['observation'])
        
        sug_title = doc.add_paragraph("【居家互動小撇步】")
        sug_title.runs[0].bold = True
        doc.add_paragraph(ai_comments['suggestion'])
        
        doc.add_paragraph("")
        
        # 5. 頁尾
        footer = doc.add_paragraph("評量代號說明： A(主動熟練)  R(表現良好)  D(發展中/需示範)  N(未觀察/需協助)")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.style = 'Quote'

    my_bar.empty()
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio

def color_grade(val):
    if val == 'A': return 'background-color: #d4edda; color: green; font-weight: bold;'
    if val == 'R': return 'background-color: #fff3cd; color: #856404; font-weight: bold;'
    if val == 'D': return 'background-color: #ffeeba; color: orange; font-weight: bold;'
    if val == 'N': return 'background-color: #f8d7da; color: red; font-weight: bold;'
    return ''

# --- 4. 主頁面邏輯 ---

if menu == "📝 批次上傳與辨識":
    st.title("📝 評量表批次處理 (v1.8)")
    st.info("💡 請上傳不同學習區的照片（例如：語文區照片+數學區照片），系統會自動識別並分類。")
    
    uploaded_files = st.file_uploader("請選擇評量表照片", type=['jpg', 'png', 'jpeg'], accept_multiple_files=True)
    
    if uploaded_files and st.button("🚀 開始分析並歸檔"):
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        all_data = [] # 用來存 DataFrame
        raw_records = [] # 用來存原始結構資料 (給 Word 生成用)

        for i, file in enumerate(uploaded_files):
            status_text.text(f"正在分析第 {i+1} 張照片...")
            image = Image.open(file)
            result = analyze_image(image)
            
            if result:
                area = result.get("area", "未知區域")
                headers = result.get("headers", ["指標1","指標2","指標3","指標4"])[:4]
                
                for s in result.get("students", []):
                    # 1. 存入 DataFrame 用的扁平資料
                    row = {"幼兒姓名": s.get("name"), "學習區": area}
                    scores = s.get("scores", [])
                    for idx, score in enumerate(scores):
                        if idx < 4: row[headers[idx]] = score
                    row["備註"] = s.get("note")
                    all_data.append(row)
                    
                    # 2. 存入結構化資料 (給 AI 寫作用)
                    details = []
                    for idx, score in enumerate(scores):
                        if idx < 4:
                            details.append({"idx": headers[idx], "score": score})
                            
                    raw_records.append({
                        "name": s.get("name"),
                        "area": area,
                        "details": details,
                        "note": s.get("note")
                    })

            progress_bar.progress((i + 1) / len(uploaded_files))
            time.sleep(1)

        if all_data:
            st.session_state['class_df'] = pd.DataFrame(all_data)
            st.session_state['raw_records'] = raw_records # 這是生成 Word 的關鍵
            st.success(f"已成功讀取 {len(uploaded_files)} 張表單，共 {len(all_data)} 筆紀錄！")

    if 'class_df' in st.session_state:
        st.divider()
        st.subheader("📊 資料預覽")
        st.dataframe(st.session_state['class_df'], use_container_width=True)

elif menu == "📄 產生整合評量報告":
    st.title("📄 整合評量報告生成中心")
    
    if 'raw_records' in st.session_state and len(st.session_state['raw_records']) > 0:
        records = st.session_state['raw_records']
        
        # 依姓名進行歸戶 (Grouping)
        grouped_data = {}
        for r in records:
            name = r['name']
            if name not in grouped_data:
                grouped_data[name] = []
            grouped_data[name].append(r)
            
        st.success(f"目前資料庫中共有 {len(grouped_data)} 位幼兒的完整學習歷程。")
        st.info("點擊下方按鈕後，AI 將會為每一位幼兒「閱讀」所有區域的成績，並撰寫客製化的觀察報告。這可能需要幾分鐘，請耐心等候。")
        
        if st.button("✨ 啟動 AI 寫作並下載報告"):
            doc_file = create_integrated_word(grouped_data)
            
            st.download_button(
                label="📥 下載全班整合報告 (Word)",
                data=doc_file,
                file_name="篤行幼兒園_全班整合評量報告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.balloons()
            
    else:
        st.warning("⚠️ 尚無資料，請先至「批次上傳」頁面分析照片。")
